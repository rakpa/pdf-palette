"""
Validate DOCX output against PDF source profiles.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field

from docx.oxml.ns import qn

from pdf_analysis import (
    CAPTION_RE,
    CODE_LINE_RE,
    MONOSPACE_HINTS,
    REFERENCE_RE,
    URL_RE,
    DocumentProfile,
    PageProfile,
    _count_code_blocks,
    _detect_headings,
    _is_monospace_font,
    _normalize_words,
)

WORD_RE = re.compile(r"\b[\w][\w\-']*\b", re.UNICODE)

TEXT_COVERAGE_MIN = 0.99
PAGE_TEXT_COVERAGE_MIN = 0.97
PAGE_COUNT_TOLERANCE = 0.02
IMAGE_COUNT_TOLERANCE = 0.0  # docx images must be >= pdf when using raster fallback


@dataclass
class PageValidation:
    index: int
    mode: str
    text_coverage: float
    passed: bool
    missing_words_sample: list[str] = field(default_factory=list)
    notes: list[str] = field(default_factory=list)


@dataclass
class ValidationReport:
    passed: bool
    text_coverage: float
    page_count_pdf: int
    page_count_docx: int
    page_count_delta_pct: float
    image_count_pdf: int
    image_count_docx: int
    heading_count_pdf: int
    heading_count_docx: int
    code_block_count_pdf: int
    code_block_count_docx: int
    caption_count_pdf: int
    caption_count_docx: int
    url_count_pdf: int
    url_count_docx: int
    reference_count_pdf: int
    reference_count_docx: int
    pages: list[PageValidation] = field(default_factory=list)
    failures: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)

    def to_dict(self) -> dict:
        return {
            "passed": self.passed,
            "text_coverage": round(self.text_coverage, 4),
            "page_count_pdf": self.page_count_pdf,
            "page_count_docx": self.page_count_docx,
            "page_count_delta_pct": round(self.page_count_delta_pct, 4),
            "image_count_pdf": self.image_count_pdf,
            "image_count_docx": self.image_count_docx,
            "heading_count_pdf": self.heading_count_pdf,
            "heading_count_docx": self.heading_count_docx,
            "code_block_count_pdf": self.code_block_count_pdf,
            "code_block_count_docx": self.code_block_count_docx,
            "caption_count_pdf": self.caption_count_pdf,
            "caption_count_docx": self.caption_count_docx,
            "url_count_pdf": self.url_count_pdf,
            "url_count_docx": self.url_count_docx,
            "reference_count_pdf": self.reference_count_pdf,
            "reference_count_docx": self.reference_count_docx,
            "failures": self.failures,
            "warnings": self.warnings,
            "pages": [
                {
                    "index": p.index,
                    "mode": p.mode,
                    "text_coverage": round(p.text_coverage, 4),
                    "passed": p.passed,
                    "missing_words_sample": p.missing_words_sample,
                    "notes": p.notes,
                }
                for p in self.pages
            ],
        }


def word_coverage(source: set[str], target: set[str]) -> float:
    if not source:
        return 1.0
    return len(source & target) / len(source)


def _element_text(element) -> str:
    return "".join(t.text or "" for t in element.iter(qn("w:t")))


def extract_docx_page_texts(document) -> list[str]:
    """Split DOCX body into per-page text buckets using section and page breaks."""
    body = document.element.body
    buckets: list[list[str]] = [[]]

    def _has_page_break(element) -> bool:
        for br in element.iter(qn("w:br")):
            if br.get(qn("w:type")) == "page":
                return True
        return False

    for child in body:
        if child.tag == qn("w:p"):
            if _has_page_break(child) and buckets[-1]:
                buckets.append([])
            buckets[-1].append(_element_text(child))
            p_pr = child.find(qn("w:pPr"))
            if p_pr is not None and p_pr.find(qn("w:sectPr")) is not None:
                buckets.append([])
        elif child.tag == qn("w:tbl"):
            buckets[-1].append(_element_text(child))

    if buckets and not any(buckets[-1]):
        buckets.pop()
    return ["\n".join(parts) for parts in buckets if parts]


def count_docx_pages(document) -> int:
    texts = extract_docx_page_texts(document)
    return max(len(texts), 1)


def count_docx_images(document) -> int:
    return sum(1 for rel in document.part.rels.values() if "image" in rel.reltype)


def _docx_heading_count(document) -> int:
    spans = []
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        size = 12.0
        for run in para.runs:
            if run.font.size:
                size = max(size, run.font.size.pt)
        from pdf_analysis import SpanInfo

        spans.append(SpanInfo(text=text, font="", size=size, bbox=(0, 0, 0, 0), flags=0))
    return _detect_headings(spans)


def _docx_code_block_count(document) -> int:
    lines: list[tuple[str, bool, str]] = []
    for para in document.paragraphs:
        text = para.text
        is_mono = False
        font_name = ""
        for run in para.runs:
            if run.font.name and _is_monospace_font(run.font.name):
                is_mono = True
                font_name = run.font.name
        lines.append((text, is_mono, font_name))
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text
                    is_mono = any(
                        run.font.name and _is_monospace_font(run.font.name) for run in para.runs
                    )
                    lines.append((text, is_mono, ""))
    return _count_code_blocks(lines)


def _count_pattern(text: str, pattern: re.Pattern[str]) -> int:
    return len(pattern.findall(text))


def validate_page_text(
    profile: PageProfile,
    docx_text: str,
    *,
    min_coverage: float = PAGE_TEXT_COVERAGE_MIN,
) -> PageValidation:
    docx_words = _normalize_words(docx_text)
    coverage = word_coverage(profile.words, docx_words)
    missing = sorted(profile.words - docx_words)[:20]
    notes: list[str] = []

    for cap in profile.caption_texts:
        cap_lower = cap.lower()
        if cap_lower not in docx_text.lower() and cap_lower[:24] not in docx_text.lower():
            notes.append(f"missing caption: {cap[:80]}")

    for url in profile.url_texts:
        if url not in docx_text:
            notes.append(f"missing url: {url[:80]}")

    passed = coverage >= min_coverage and not notes
    return PageValidation(
        index=profile.index,
        mode="vector",
        text_coverage=coverage,
        passed=passed,
        missing_words_sample=missing,
        notes=notes,
    )


def validate_document(
    pdf_profile: DocumentProfile,
    document,
    *,
    page_modes: list[str] | None = None,
) -> ValidationReport:
    page_texts = extract_docx_page_texts(document)
    full_text = "\n".join(page_texts)

    docx_words = _normalize_words(full_text)
    covered_words: set[str] = set()
    for i, profile in enumerate(pdf_profile.pages):
        mode = page_modes[i] if page_modes and i < len(page_modes) else "vector"
        if mode == "image":
            covered_words |= profile.words
        else:
            page_text = page_texts[i] if i < len(page_texts) else ""
            covered_words |= profile.words & _normalize_words(page_text)
    text_cov = (
        len(covered_words) / len(pdf_profile.all_words) if pdf_profile.all_words else 1.0
    )

    page_count_docx = count_docx_pages(document)
    page_delta = abs(page_count_docx - pdf_profile.page_count) / max(pdf_profile.page_count, 1)

    image_docx = count_docx_images(document)
    heading_docx = _docx_heading_count(document)
    code_docx = _docx_code_block_count(document)
    caption_docx = _count_pattern(full_text, CAPTION_RE)
    url_docx = _count_pattern(full_text, URL_RE)
    ref_docx = _count_pattern(full_text, REFERENCE_RE)

    vector_caption_pdf = 0
    vector_url_pdf = 0
    vector_ref_pdf = 0
    for i, p in enumerate(pdf_profile.pages):
        mode = page_modes[i] if page_modes and i < len(page_modes) else "vector"
        if mode == "image":
            continue
        vector_caption_pdf += len(p.caption_texts)
        vector_url_pdf += len(p.url_texts)
        vector_ref_pdf += len(p.reference_texts)

    pages: list[PageValidation] = []
    for i, profile in enumerate(pdf_profile.pages):
        mode = (page_modes[i] if page_modes and i < len(page_modes) else "vector")
        if mode == "image":
            pages.append(
                PageValidation(
                    index=i,
                    mode="image",
                    text_coverage=1.0,
                    passed=True,
                    notes=["rendered as high-resolution page image"],
                )
            )
            continue
        docx_page_text = page_texts[i] if i < len(page_texts) else ""
        pages.append(validate_page_text(profile, docx_page_text))

    failures: list[str] = []
    warnings: list[str] = []

    if text_cov < TEXT_COVERAGE_MIN:
        failures.append(f"text coverage {text_cov:.2%} < {TEXT_COVERAGE_MIN:.0%}")
    if page_delta > PAGE_COUNT_TOLERANCE:
        failures.append(
            f"page count delta {page_delta:.2%} exceeds {PAGE_COUNT_TOLERANCE:.0%} "
            f"(pdf={pdf_profile.page_count}, docx={page_count_docx})"
        )
    if image_docx < pdf_profile.image_count and not any(
        m == "image" for m in (page_modes or [])
    ):
        failures.append(
            f"image count docx={image_docx} < pdf={pdf_profile.image_count}"
        )
    if heading_docx < pdf_profile.heading_count:
        warnings.append(
            f"heading count docx={heading_docx} < pdf={pdf_profile.heading_count}"
        )
    if code_docx < pdf_profile.code_block_count:
        warnings.append(
            f"code block count docx={code_docx} < pdf={pdf_profile.code_block_count}"
        )
    if caption_docx < vector_caption_pdf:
        failures.append(
            f"caption count docx={caption_docx} < pdf vector pages={vector_caption_pdf}"
        )
    if url_docx < vector_url_pdf:
        failures.append(f"url count docx={url_docx} < pdf vector pages={vector_url_pdf}")
    if ref_docx < vector_ref_pdf:
        warnings.append(
            f"reference count docx={ref_docx} < pdf vector pages={vector_ref_pdf}"
        )

    for pv in pages:
        if pv.mode != "image" and not pv.passed:
            failures.append(
                f"page {pv.index + 1} text coverage {pv.text_coverage:.2%}"
            )

    return ValidationReport(
        passed=len(failures) == 0,
        text_coverage=text_cov,
        page_count_pdf=pdf_profile.page_count,
        page_count_docx=page_count_docx,
        page_count_delta_pct=page_delta,
        image_count_pdf=pdf_profile.image_count,
        image_count_docx=image_docx,
        heading_count_pdf=pdf_profile.heading_count,
        heading_count_docx=heading_docx,
        code_block_count_pdf=pdf_profile.code_block_count,
        code_block_count_docx=code_docx,
        caption_count_pdf=pdf_profile.caption_count,
        caption_count_docx=caption_docx,
        url_count_pdf=pdf_profile.url_count,
        url_count_docx=url_docx,
        reference_count_pdf=pdf_profile.reference_count,
        reference_count_docx=ref_docx,
        pages=pages,
        failures=failures,
        warnings=warnings,
    )
