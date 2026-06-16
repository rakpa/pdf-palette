#!/usr/bin/env python3
"""
Pixel-accurate PDF → DOCX conversion (local, no cloud APIs).

Pipeline:
  1. Profile PDF pages from object coordinates (text, images, headings, code).
  2. Reconstruct each page in isolation via pdf2docx (coordinate-based).
  3. Validate per-page text coverage; fall back to high-DPI page image if needed.
  4. Compose pages without cross-page content merge.
  5. Apply fidelity cleanup (preserve page breaks) + monospace/code styling.
  6. Validate output against source thresholds.
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

from docx_page_builder import build_fidelity_docx
from docx_validator import validate_document
from pdf_analysis import analyze_pdf


def log(msg: str) -> None:
    print(msg, file=sys.stderr, flush=True)


def run_ocr_if_needed(pdf_path: Path, likely_scanned: bool) -> Path:
    if not likely_scanned:
        return pdf_path

    try:
        import ocrmypdf
    except ImportError:
        log("ocrmypdf not installed — continuing without OCR (scanned PDF quality may be lower)")
        return pdf_path

    ocr_path = pdf_path.with_name(f"{pdf_path.stem}_ocr.pdf")
    log("running OCR for scanned PDF…")
    try:
        ocrmypdf.ocr(
            str(pdf_path),
            str(ocr_path),
            skip_text=True,
            optimize=0,
            output_type="pdf",
            progress_bar=False,
        )
        return ocr_path
    except Exception as exc:
        log(f"OCR skipped: {exc}")
        if ocr_path.exists():
            ocr_path.unlink(missing_ok=True)
        return pdf_path


def fidelity_cleanup(docx_path: Path) -> dict:
    from docx import Document

    from docx_cleanup import fidelity_docx_cleanup

    doc = Document(str(docx_path))
    stats = fidelity_docx_cleanup(doc)
    doc.save(str(docx_path))
    if any(stats.values()):
        log(f"fidelity cleanup: {stats}")
    return stats


def apply_monospace_styles(docx_path: Path, snippets: list[str], code_lines: list[str]) -> int:
    if not snippets and not code_lines:
        return 0

    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(str(docx_path))
    mono_font = "Consolas"
    styled = 0
    unique = sorted(
        {s for s in snippets + code_lines if s and len(s) >= 2},
        key=len,
        reverse=True,
    )

    def style_paragraph(para) -> None:
        nonlocal styled
        para_text = para.text
        if not para_text.strip():
            return
        for snippet in unique:
            if snippet not in para_text:
                continue
            for run in para.runs:
                if snippet in run.text or run.text in snippet:
                    run.font.name = mono_font
                    r = run._element.get_or_add_rPr()
                    r.rFonts.set(qn("w:ascii"), mono_font)
                    r.rFonts.set(qn("w:hAnsi"), mono_font)
                    r.rFonts.set(qn("w:cs"), mono_font)
                    r.rFonts.set(qn("w:eastAsia"), mono_font)
                    styled += 1

    for para in doc.paragraphs:
        style_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    style_paragraph(para)

    doc.save(str(docx_path))
    return styled


def convert(input_pdf: str, output_docx: str) -> dict:
    src = Path(input_pdf).resolve()
    dst = Path(output_docx).resolve()
    dst.parent.mkdir(parents=True, exist_ok=True)

    if not src.exists():
        raise FileNotFoundError(f"Input PDF not found: {src}")

    profile = analyze_pdf(src)
    agg = profile.aggregate()
    log(
        f"PDF profile: {agg['page_count']} pages, {agg['text_chars']} chars, "
        f"{agg['image_count']} images, {agg['heading_count']} headings, "
        f"{agg['code_block_count']} code blocks"
    )

    work_pdf = run_ocr_if_needed(src, profile.likely_scanned)
    ocr_used = work_pdf != src

    try:
        if work_pdf != src:
            profile = analyze_pdf(work_pdf)

        build_stats = build_fidelity_docx(work_pdf, dst, profile, log)
        cleanup_stats = fidelity_cleanup(dst)

        mono_snippets = list(
            set(profile.monospace_snippets + profile.code_lines)
        )
        styled_runs = apply_monospace_styles(dst, mono_snippets, profile.code_lines)

        from docx import Document

        doc = Document(str(dst))
        validation = validate_document(
            profile,
            doc,
            page_modes=build_stats.get("page_modes"),
        )

        if not validation.passed:
            log(f"validation warnings/failures: {validation.failures}")
            if validation.warnings:
                log(f"validation notes: {validation.warnings}")

        table_count = len(doc.tables)
        image_count = sum(
            1 for rel in doc.part.rels.values() if "image" in rel.reltype
        )

        return {
            "ok": True,
            "output": str(dst),
            "page_count": profile.page_count,
            "byte_length": dst.stat().st_size,
            "pdf_image_count": profile.image_count,
            "docx_image_count": image_count,
            "docx_table_count": table_count,
            "monospace_runs_styled": styled_runs,
            "ocr_used": ocr_used,
            "engine": "pdf2docx-fidelity",
            "build": build_stats,
            "cleanup": cleanup_stats,
            "validation": validation.to_dict(),
        }
    finally:
        if ocr_used and work_pdf.exists():
            work_pdf.unlink(missing_ok=True)


def main() -> int:
    if len(sys.argv) != 3:
        print(json.dumps({"ok": False, "error": "Usage: pdf_to_docx.py <input.pdf> <output.docx>"}))
        return 1

    try:
        result = convert(sys.argv[1], sys.argv[2])
        print(json.dumps(result))
        return 0
    except Exception as exc:
        print(json.dumps({"ok": False, "error": str(exc)}))
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
