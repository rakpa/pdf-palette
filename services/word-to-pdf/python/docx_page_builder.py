"""
Per-page DOCX construction with coordinate reconstruction and image fallback.

Each PDF page is converted in isolation so cover content never merges with
later pages. Pages that fail text-coverage validation are replaced with
high-resolution embedded renders of the PDF page region.
"""

from __future__ import annotations

import io
import tempfile
from pathlib import Path
from typing import Callable

from docx import Document
from docx.shared import Pt

from docx_validator import PAGE_TEXT_COVERAGE_MIN, validate_page_text
from pdf_analysis import DocumentProfile, PageProfile

LogFn = Callable[[str], None]

DEFAULT_RENDER_DPI = 300


def convert_page_vector(pdf_path: Path, page_index: int, output_path: Path) -> None:
    import fitz
    from pdf2docx import Converter

    src = fitz.open(pdf_path)
    single = fitz.open()
    single.insert_pdf(src, from_page=page_index, to_page=page_index)
    src.close()

    temp_pdf = output_path.with_suffix(".single.pdf")
    single.save(str(temp_pdf))
    single.close()

    converter = Converter(str(temp_pdf))
    try:
        converter.convert(str(output_path))
    finally:
        converter.close()
        temp_pdf.unlink(missing_ok=True)


def build_page_image_docx(
    pdf_path: Path,
    page_index: int,
    output_path: Path,
    profile: PageProfile,
    *,
    dpi: int = DEFAULT_RENDER_DPI,
) -> None:
    import fitz

    doc = fitz.open(pdf_path)
    page = doc[page_index]
    zoom = dpi / 72.0
    pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
    doc.close()

    word_doc = Document()
    section = word_doc.sections[0]
    section.page_width = Pt(profile.width)
    section.page_height = Pt(profile.height)
    section.left_margin = Pt(0)
    section.right_margin = Pt(0)
    section.top_margin = Pt(0)
    section.bottom_margin = Pt(0)

    stream = io.BytesIO(pix.tobytes("png"))
    para = word_doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run()
    run.add_picture(stream, width=Pt(profile.width), height=Pt(profile.height))
    word_doc.save(str(output_path))


def compose_page_docx_files(page_paths: list[Path], output_path: Path) -> None:
    if not page_paths:
        raise RuntimeError("No page DOCX files to compose.")
    try:
        from docxcompose.composer import Composer
    except ImportError as exc:
        raise RuntimeError(
            "docxcompose is required. Run: pip install -r services/word-to-pdf/python/requirements.txt"
        ) from exc

    from docx.enum.text import WD_BREAK

    master = Document(str(page_paths[0]))
    composer = Composer(master)
    for path in page_paths[1:]:
        # Hard page break so each PDF page stays isolated in the output.
        break_para = master.add_paragraph()
        run = break_para.add_run()
        run.add_break(WD_BREAK.PAGE)
        composer.append(Document(str(path)))
    composer.save(str(output_path))


def build_fidelity_docx(
    pdf_path: Path,
    output_path: Path,
    profile: DocumentProfile,
    log: LogFn,
    *,
    render_dpi: int = DEFAULT_RENDER_DPI,
) -> dict:
    """
    Convert PDF to DOCX page-by-page with validation and image fallback.

    Returns build statistics including per-page modes.
    """
    page_modes: list[str] = []
    page_coverages: list[float] = []
    image_fallback_pages: list[int] = []

    with tempfile.TemporaryDirectory(prefix="pdf-palette-pages-") as tmp:
        tmp_dir = Path(tmp)
        page_paths: list[Path] = []

        for page_profile in profile.pages:
            idx = page_profile.index
            vector_path = tmp_dir / f"page_{idx:04d}_vector.docx"
            image_path = tmp_dir / f"page_{idx:04d}_image.docx"
            final_path = tmp_dir / f"page_{idx:04d}.docx"

            log(f"page {idx + 1}/{profile.page_count}: coordinate reconstruction…")
            try:
                convert_page_vector(pdf_path, idx, vector_path)
                from docx import Document as DocxDocument

                vec_doc = DocxDocument(str(vector_path))
                page_texts = _single_page_text(vec_doc)
                docx_text = page_texts[0] if page_texts else ""
                validation = validate_page_text(page_profile, docx_text)
                page_coverages.append(validation.text_coverage)

                if validation.passed:
                    page_modes.append("vector")
                    final_path.write_bytes(vector_path.read_bytes())
                    log(
                        f"page {idx + 1}: vector OK "
                        f"(coverage {validation.text_coverage:.1%})"
                    )
                else:
                    raise ValueError(
                        f"text coverage {validation.text_coverage:.2%} "
                        f"< {PAGE_TEXT_COVERAGE_MIN:.0%}"
                    )
            except Exception as exc:
                log(
                    f"page {idx + 1}: vector reconstruction insufficient "
                    f"({exc}) — embedding {render_dpi} DPI page image"
                )
                build_page_image_docx(
                    pdf_path, idx, image_path, page_profile, dpi=render_dpi
                )
                page_modes.append("image")
                image_fallback_pages.append(idx)
                page_coverages.append(1.0)
                final_path.write_bytes(image_path.read_bytes())

            page_paths.append(final_path)

        log(f"composing {len(page_paths)} isolated pages…")
        compose_page_docx_files(page_paths, output_path)

    return {
        "page_modes": page_modes,
        "page_coverages": [round(c, 4) for c in page_coverages],
        "image_fallback_pages": image_fallback_pages,
        "vector_pages": sum(1 for m in page_modes if m == "vector"),
        "image_pages": sum(1 for m in page_modes if m == "image"),
        "render_dpi": render_dpi,
    }


def _single_page_text(document) -> list[str]:
    from docx_validator import extract_docx_page_texts

    return extract_docx_page_texts(document)
